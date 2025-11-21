import os
import json
from docx import Document
import re

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SUBMISSIONS_FOLDER = os.path.join(BASE_DIR, "submissions")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "extracted")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Regex patterns
email_pattern = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", re.I)
table_pattern = re.compile(r"^Table\s*\d+", re.I)
figure_pattern = re.compile(r"^Fig(?:ure)?\s*\d+", re.I)

# Keywords for automatic section detection
section_keywords = {
    "Abstract": ["abstract", "abstrak"],
    "Introduction": ["introduction"],
    "Materials and Methods": ["materials and methods", "methodology", "methods"],
    "Results and Discussion": ["results", "discussion", "results and discussion"],
    "Conclusions": ["conclusion", "concluding remarks"],
    "References": ["references", "bibliography", "daftar pustaka"]
}

def label_paragraph(text, style):
    lower = text.lower()
    for label, keywords in section_keywords.items():
        for kw in keywords:
            if kw.lower() in lower:
                return label

    
    # Check for email
    if email_pattern.search(text):
        return "Corresponding Email"
    
    # Check for Table
    if table_pattern.match(text):
        return "Table Caption"
    
    # Check for Figure
    if figure_pattern.match(text):
        return "Figure Caption"
    
    # Check for keywords
    for label, keywords in section_keywords.items():
        for kw in keywords:
            if kw in lower:
                return label
    
    # Style-based guesses
    if style.startswith("Heading"):
        return "Section Heading"
    
    # Default unknown
    return "Unknown"

def extract_docx(file_path):
    doc = Document(file_path)
    content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name
            label = label_paragraph(text, style)
            content.append({
                "text": text,
                "style": style,
                "label": label
            })

    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        tables.append({
            "content": table_data,
            "label": "Table"
        })

    return {"paragraphs": content, "tables": tables}

# Process all submissions
for file_name in os.listdir(SUBMISSIONS_FOLDER):
    # skip temporary word file
    if file_name.startswith("~$"):
        print("Skipping temporary file")
        continue

    file_path = os.path.join(SUBMISSIONS_FOLDER, file_name)
    
    if file_name.endswith(".docx"):
        data = extract_docx(file_path)
    elif file_name.endswith(".doc"):
        print(f"Skipping .doc file for now: {file_name}")
        continue
    else:
        print(f"Unsupported file type: {file_name}")
        continue

    # Save extracted and pre-labeled data as JSON
    output_file = os.path.join(OUTPUT_FOLDER, file_name + ".json")
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"{file_name} extracted → {len(data['paragraphs'])} paragraphs, {len(data['tables'])} tables (pre-labeled)")
