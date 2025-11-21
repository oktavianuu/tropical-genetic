import os
import json
from docx import Document

# Paths
SUBMISSIONS_FOLDER = "submissions"
OUTPUT_FOLDER = "extracted"

# Create output folder if not exists
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def extract_docx(file_path):
    doc = Document(file_path)
    content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name  # 'Heading 1', 'Normal', etc.
            content.append({
                "text": text,
                "style": style
            })

    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        tables.append(table_data)

    return {"paragraphs": content, "tables": tables}

# Process all submissions
for file_name in os.listdir(SUBMISSIONS_FOLDER):
    file_path = os.path.join(SUBMISSIONS_FOLDER, file_name)
    
    if file_name.endswith(".docx"):
        data = extract_docx(file_path)
    elif file_name.endswith(".doc"):
        print(f"Skipping .doc file for now: {file_name}")
        continue
    else:
        print(f"Unsupported file type: {file_name}")
        continue

    # Save extracted data as JSON
    output_file = os.path.join(OUTPUT_FOLDER, file_name + ".json")
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"{file_name} extracted → {len(data['paragraphs'])} paragraphs, {len(data['tables'])} tables")
